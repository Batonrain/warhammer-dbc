# tools/book-coverage.py
# ════════════════════════════════════════════════════════════════════════
#  Сколько текста исходника книги реально доехало до
#  packs-src/books/<slug>.json.
#
#  Метод — покрытие словесными шинглами (6 слов). Устойчиво к склейке
#  колонок и перестановке абзацев: если текст есть где угодно в книге,
#  шингл найдётся. Не найден — текста нет, это дыра, а не перестановка.
#
#  Только читает. Ничего не правит.
#
#    python tools/book-coverage.py [slug ...] [--holes=N] [--dump=ДИР]
# ════════════════════════════════════════════════════════════════════════
import sys, os, re, json, zipfile, io
from html import unescape

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
SRC = r"D:\tRPG\Warhammer\Самоделки"
DL = r"C:\Users\Derbius\Downloads"

SOURCES = {
    "aeldari-branches": (SRC, "Книга Аэльдари_ Ответвления.zip"),
    "aeldari":          (SRC, "Книга Аэльдари.zip"),
    "battles":          (SRC, "Книга Битв.zip"),
    "chaos":            (SRC, "DoomBC_S_Chaos.pdf"),
    "core":             (SRC, "DoomBC_Core .pdf"),
    "daemonic-shells":  (SRC, "Книга Демонических Оболочек.pdf"),
    "diseases":         (SRC, "Книга Болезней.zip"),
    "divinations-book": (SRC, "Родные миры и Предсказания.zip"),
    "eldar-vehicles":   (SRC, "Книга Эльдар_ Техника.zip"),
    "machines":         (SRC, "DoomBC_Machines.pdf"),
    "necrons":          (SRC, "Книга Некрон.zip"),
    "origins-book":     (SRC, "Родные миры и Предсказания.zip"),
    "power-armour":     (SRC, "Силовая броня_ без шлема и особенности.zip"),
    "toad-psykers":     (SRC, "Жабья Книга Псайкеров.pdf"),
    "tyranids":         (SRC, "Тираниды DBC.zip"),
    "void":             (DL,  "Книга Пустоты v.2 (1).docx"),
}

CORPUS = []
N = 6                      # длина шингла в словах
HYPHEN = re.compile(r"(\w)-\s*\n\s*(\w)")
WORD = re.compile(r"[0-9A-Za-zА-Яа-яЁё]+")


def words(text):
    return [w.lower().replace("ё", "е") for w in WORD.findall(text or "")]


def strip_html(html):
    h = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", html or "")
    h = re.sub(r"(?s)<[^>]+>", " ", h)
    # Экспорт Google Docs пишет кириллицу числовыми сущностями (&#1057;).
    # Без unescape весь русский текст исходника читается как цифры — замер
    # тогда врёт про потерю всей книги.
    return unescape(h)


# ── источники ───────────────────────────────────────────────────────────
def src_pdf(path):
    """[(метка, текст)] по страницам PDF."""
    import pymupdf
    doc = pymupdf.open(path)
    out = []
    for i, page in enumerate(doc, 1):
        # Узкие колонки книги переносят слова через дефис в конце строки.
        # Без склейки половинки читаются как два отдельных слова, и любой
        # абзац с переносом выглядит потерянным, хотя в JSON он целый.
        text = HYPHEN.sub(chr(92) + '1' + chr(92) + '2', page.get_text('text'))
        out.append((f"стр.{i}", text))
    doc.close()
    return out


def src_zip(path):
    """[(метка, текст)] — экспорт Google Docs, один index.html на всё."""
    out = []
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if not name.lower().endswith((".html", ".htm")):
                continue
            raw = z.read(name).decode("utf-8", "replace")
            out.append((name, strip_html(raw)))
    return out


def src_docx(path):
    """[(метка, текст)] — абзацы и таблицы docx одним куском."""
    import docx
    d = docx.Document(path)
    buf = []
    for p in d.paragraphs:
        buf.append(p.text)
    # Объединённая ячейка возвращается python-docx на каждую покрытую
    # позицию — без отсева её текст размножается и даёт ложные «дыры».
    for t in d.tables:
        seen = set()
        for row in t.rows:
            for c in row.cells:
                key = id(c._tc)
                if key in seen:
                    continue
                seen.add(key)
                buf.append(c.text)
    return [(os.path.basename(path), "\n".join(buf))]


def load_source(slug):
    folder, name = SOURCES[slug]
    path = os.path.join(folder, name)
    if not os.path.exists(path):
        return None, path
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return src_pdf(path), path
    if ext == ".zip":
        return src_zip(path), path
    if ext == ".docx":
        return src_docx(path), path
    raise SystemExit(f"неизвестный тип источника: {path}")


# ── цель ────────────────────────────────────────────────────────────────
def load_target(slug):
    p = os.path.join(ROOT, "packs-src", "books", slug + ".json")
    book = json.load(open(p, encoding="utf-8"))
    buf = []
    for e in book.get("entries", []):
        buf.append(e.get("name") or "")
        for pg in e.get("pages", []):
            buf.append(pg.get("name") or "")
            buf.append(strip_html(pg.get("html")))
    return book, "\n".join(buf)


def json_text(node, buf):
    """Все строки JSON-документа подряд — имена, описания, html карточек."""
    if isinstance(node, str):
        buf.append(node)
    elif isinstance(node, list):
        for v in node:
            json_text(v, buf)
    elif isinstance(node, dict):
        for v in node.values():
            json_text(v, buf)


def corpus_all():
    """Текст ВСЕГО packs-src: книга могла отдать таблицу в карточки предметов,
    и тогда текст не потерян, а просто лежит не в журнале."""
    buf = []
    for dirpath, _dirs, files in os.walk(os.path.join(ROOT, "packs-src")):
        for fn in files:
            if fn.endswith(".json"):
                try:
                    json_text(json.load(open(os.path.join(dirpath, fn), encoding="utf-8")), buf)
                except Exception:
                    pass
    return strip_html(chr(10).join(buf))


def shingles(ws):
    return {tuple(ws[i:i + N]) for i in range(len(ws) - N + 1)}


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    opts = {a.split("=")[0]: (a.split("=", 1)[1] if "=" in a else True)
            for a in sys.argv[1:] if a.startswith("--")}
    hole_min = int(opts.get("--holes", 20))
    dump = opts.get("--dump")
    slugs = args or sorted(SOURCES)

    if dump:
        os.makedirs(dump, exist_ok=True)
    if opts.get("--corpus") == "all":
        CORPUS.append(corpus_all())

    print(f"{'книга':<18} {'слов ист.':>9} {'слов JSON':>9} {'покрытие':>9} {'дыр':>5} {'слов в дырах':>12}")
    print("-" * 70)
    for slug in slugs:
        chunks, path = load_source(slug)
        if chunks is None:
            print(f"{slug:<18} ИСХОДНИК НЕ НАЙДЕН: {path}")
            continue
        book, tgt_text = load_target(slug)
        if opts.get("--corpus") == "all":
            tgt_text = tgt_text + chr(10) + CORPUS[0]
        tgt = shingles(words(tgt_text))

        rows, total_src, covered, holes, hole_words = [], 0, 0, [], 0
        for label, text in chunks:
            ws = words(text)
            total_src += len(ws)
            if len(ws) < N:
                continue
            mark = [False] * len(ws)
            for i in range(len(ws) - N + 1):
                if tuple(ws[i:i + N]) in tgt:
                    for j in range(i, i + N):
                        mark[j] = True
            covered += sum(mark)
            i = 0
            while i < len(mark):
                if mark[i]:
                    i += 1
                    continue
                j = i
                while j < len(mark) and not mark[j]:
                    j += 1
                if j - i >= hole_min:
                    holes.append((label, j - i, " ".join(ws[i:j])))
                    hole_words += j - i
                i = j
        pct = 100.0 * covered / total_src if total_src else 0.0
        tgt_words = len(words(tgt_text))
        print(f"{slug:<18} {total_src:>9} {tgt_words:>9} {pct:>8.1f}% {len(holes):>5} {hole_words:>12}")
        if dump:
            with open(os.path.join(dump, slug + ".txt"), "w", encoding="utf-8") as f:
                f.write(f"# {slug} — исходник {path}\n")
                f.write(f"# покрытие {pct:.1f}%  дыр {len(holes)}  слов в дырах {hole_words}\n\n")
                for label, n, txt in sorted(holes, key=lambda h: -h[1]):
                    f.write(f"--- {label}  ({n} слов)\n{txt}\n\n")


if __name__ == "__main__":
    main()
