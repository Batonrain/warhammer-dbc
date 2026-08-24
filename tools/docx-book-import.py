# tools/docx-book-import.py
# ════════════════════════════════════════════════════════════════════════
#  Конвертер .docx (Word/Google Docs) в packs-src/books/<slug>.json.
#
#  Сестра tools/html-book-import.py: тот же принцип (заголовки Word ->
#  структура книги, жирность/цвет из форматирования рана, а не из пикселей
#  картинки), но источник — сам .docx через python-docx, без промежуточного
#  HTML-экспорта. Body документа обходится ПО ПОРЯДКУ (параграфы и таблицы
#  вперемешку) через XML, не через document.paragraphs/.tables отдельно —
#  иначе порядок между текстом и таблицами теряется.
#
#  Заголовки: Heading 1 -> глава (entries[]), Heading 2..6 -> раздел
#  (pages[], поле level = уровень-1, как у book-outline.py для PDF-закладок).
#
#  Использование:
#    python tools/docx-book-import.py <файл.docx> <slug> --title="..." --file="..."
#    [--out=packs-src/books/<slug>.json]
# ════════════════════════════════════════════════════════════════════════
import sys, os, json, re
import docx
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
HEADING_RE = re.compile(r"^Heading (\d)$")


def run_style(run):
    info = {}
    if run.bold:
        info["bold"] = True
    if run.italic:
        info["italic"] = True
    color = run.font.color
    if color is not None and color.type is not None and color.rgb is not None:
        rgb = str(color.rgb)
        if rgb.upper() != "000000":
            info["color"] = "#" + rgb.lower()
    return info


def _escape(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _styled(text, run):
    if not text:
        return ""
    text = _escape(text)
    style = run_style(run)
    if style.get("bold"):
        text = f"<strong>{text}</strong>"
    if style.get("italic"):
        text = f"<em>{text}</em>"
    if style.get("color"):
        text = f'<span style="color:{style["color"]}">{text}</span>'
    return text


def para_inline_html(paragraph):
    """Обходит XML параграфа напрямую, не через .runs — иначе текст внутри
    <w:hyperlink> (гиперссылки) пропадает молча, это отдельный элемент,
    не обычный <w:r>."""
    out = []
    part = paragraph.part
    for child in paragraph._p.iterchildren():
        if child.tag == qn("w:r"):
            run = docx.text.run.Run(child, paragraph)
            out.append(_styled(run.text, run))
        elif child.tag == qn("w:hyperlink"):
            rid = child.get(qn("r:id"))
            href = None
            if rid:
                try:
                    href = part.rels[rid].target_ref
                except KeyError:
                    href = None
            inner = []
            for sub in child.iterchildren():
                if sub.tag == qn("w:r"):
                    run = docx.text.run.Run(sub, paragraph)
                    inner.append(_styled(run.text, run))
            text = "".join(inner)
            if text:
                # href — атрибут: помимо &/<> надо экранировать и кавычку,
                # иначе URL с «"» разорвёт атрибут (та же причина, что _escape
                # для текста ранов).
                href_esc = _escape(href).replace('"', "&quot;") if href else None
                out.append(f'<a href="{href_esc}">{text}</a>' if href else text)
    return "".join(out)


def para_plain_text(paragraph):
    return paragraph.text.strip()


def table_html(table):
    """Как convert_block() таблиц в html-book-import.py: colspan/rowspan/th.

    По XML, не по row.cells: python-docx для смёрженной ячейки отдаёт один и
    тот же _tc во всех «накрытых» позициях, из-за чего старый обход с
    seen_in_row терял и colspan, и rowspan — содержимое повторялось в каждой
    строке вертикального слияния, а горизонтальное схлопывалось в одну <td>
    без атрибута. В XML же всё явно: w:gridSpan — colspan, w:vMerge
    restart/continue — начало/продолжение вертикального слияния,
    w:tblHeader на строке — строка-шапка (<th>).
    """
    rows_html = []
    trs = table._tbl.tr_lst
    for ri, tr in enumerate(trs):
        cells_html = []
        is_header = tr.trPr is not None and tr.trPr.find(qn("w:tblHeader")) is not None
        grid_col = 0
        for tc in tr.tc_lst:
            span = tc.grid_span
            if tc.vMerge is not None and tc.vMerge != "restart":
                # продолжение вертикального слияния — ячейка уже отрисована
                # выше со своим rowspan
                grid_col += span
                continue
            rowspan = 1
            if tc.vMerge == "restart":
                for tr2 in trs[ri + 1:]:
                    gc2, cont = 0, False
                    for tc2 in tr2.tc_lst:
                        if gc2 == grid_col:
                            cont = tc2.vMerge is not None and tc2.vMerge != "restart"
                            break
                        gc2 += tc2.grid_span
                    if not cont:
                        break
                    rowspan += 1
            tag = "th" if is_header else "td"
            attrs = ""
            if span > 1:
                attrs += f' colspan="{span}"'
            if rowspan > 1:
                attrs += f' rowspan="{rowspan}"'
            cell = docx.table._Cell(tc, table)
            parts = []
            for para in cell.paragraphs:
                t = para_inline_html(para).strip()
                if t:
                    parts.append(f"<p>{t}</p>" if len(cell.paragraphs) > 1 else t)
            cells_html.append(f"<{tag}{attrs}>{''.join(parts)}</{tag}>")
            grid_col += span
        if cells_html:
            rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
    if not rows_html:
        return ""
    return f"<table>{''.join(rows_html)}</table>"


def iter_body_blocks(document):
    """Параграфы и таблицы вперемешку, в реальном порядке документа."""
    body = document.element.body
    paragraphs_by_elem = {p._p: p for p in document.paragraphs}
    tables_by_elem = {t._tbl: t for t in document.tables}
    for child in body.iterchildren():
        if child.tag == qn("w:p") and child in paragraphs_by_elem:
            yield ("p", paragraphs_by_elem[child])
        elif child.tag == qn("w:tbl") and child in tables_by_elem:
            yield ("tbl", tables_by_elem[child])


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    opts = dict((a[2:].split("=", 1) + [True])[:2] for a in sys.argv[1:] if a.startswith("--"))
    docx_path, slug = args[0], args[1]
    title = opts.get("title", slug)
    file_field = opts.get("file", os.path.basename(docx_path))
    out_path = opts.get("out", f"packs-src/books/{slug}.json")

    document = docx.Document(docx_path)

    entries = []
    cur_chapter = None
    cur_page = None
    page_counter = 0

    def flush_page():
        nonlocal cur_page
        if cur_page and cur_page["html"].strip():
            page = {
                "name": cur_page["name"],
                "pdfPage": cur_page["idx"],
                "html": cur_page["html"],
            }
            if cur_page["level"] > 1:
                page["level"] = cur_page["level"]
            cur_chapter["pages"].append(page)
        cur_page = None

    for kind, block in iter_body_blocks(document):
        if kind == "p":
            style_name = block.style.name if block.style else "Normal"
            m = HEADING_RE.match(style_name or "")
            if m:
                level_num = int(m.group(1))
                name = para_plain_text(block)
                if not name:
                    continue
                if level_num == 1:
                    flush_page()
                    cur_chapter = {"name": name, "pages": []}
                    entries.append(cur_chapter)
                    continue
                flush_page()
                if cur_chapter is None:
                    cur_chapter = {"name": title, "pages": []}
                    entries.append(cur_chapter)
                page_counter += 1
                level = level_num - 1
                htag = f"h{min(level + 1, 6)}"
                inner = para_inline_html(block).strip() or name
                cur_page = {
                    "name": name, "idx": page_counter, "level": level,
                    "html": f"<{htag}>{inner}</{htag}>",
                }
                continue
            text = para_inline_html(block).strip()
            if not text:
                continue
            if style_name == "Subtitle":
                # жирность стиля Subtitle задана на уровне стиля параграфа, а не
                # рана — run.bold её не видит, поэтому оборачиваем вручную здесь
                text = f"<strong>{text}</strong>"
            frag = f"<p>{text}</p>"
        else:
            frag = table_html(block)
            if not frag:
                continue
        if cur_page is None:
            if cur_chapter is None:
                cur_chapter = {"name": title, "pages": []}
                entries.append(cur_chapter)
            page_counter += 1
            cur_page = {"name": cur_chapter["name"], "idx": page_counter, "level": 1, "html": ""}
        cur_page["html"] += frag
    flush_page()

    data = {
        "slug": slug,
        "title": title,
        "file": file_field,
        "pdfPages": page_counter,
        "entries": entries,
    }
    out_full = os.path.join(ROOT, out_path)
    json.dump(data, open(out_full, "w", encoding="utf-8", newline="\n"), ensure_ascii=False, indent=1)

    total_pages = sum(len(e["pages"]) for e in entries)
    print(f"Готово: {slug} — {len(entries)} глав, {total_pages} разделов -> {out_path}")


if __name__ == "__main__":
    main()
